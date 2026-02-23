"""
Servicio para generación de documentos DOCX y PDF a partir de plantillas ISO.
Usa las plantillas Word originales en instance/apps/helpdesk/templates/ y las llena
con datos de los tickets, preservando el formato exacto del documento certificado.

Para PDF: convierte el DOCX generado usando LibreOffice CLI.
"""
import os
import shutil
import subprocess
import tempfile
import zipfile
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.opc.part import Part as OpcPart
from docx.opc.packuri import PackURI
from flask import current_app
from PIL import Image, ImageDraw

from itcj.apps.helpdesk.models.ticket import Ticket
import logging

logger = logging.getLogger(__name__)

# Rutas relativas a instance/
TEMPLATE_SOLICITUD = 'apps/helpdesk/templates/FORMATO PARA SOLICITUD DE MANTENIMIENTO DE EQUIPO INFORM_TICO(2).docx'
TEMPLATE_ORDEN = 'apps/helpdesk/templates/FORMATO PARA ORDEN DE TRABAJO DE MANTENIMIENTO DE EQUIPO INFORM_TICO(2).docx'

# Constantes de auto-sizing
DEFAULT_FONT_SIZE_PT = 11
MIN_FONT_SIZE_PT = 7
EMU_PER_PT = 12700


def _get_template_path(template_name: str) -> str:
    """Obtiene la ruta absoluta de una plantilla en instance/."""
    instance_path = current_app.instance_path
    path = os.path.join(instance_path, template_name)
    if not os.path.exists(path):
        raise FileNotFoundError(f'Plantilla no encontrada: {path}')
    return path


def _format_date(dt) -> str:
    """Formatea datetime a DD/MM/AAAA."""
    return dt.strftime('%d/%m/%Y') if dt else ''


# ==================== AUTO-SIZING ====================

def _calculate_font_size(text: str, available_height_emu: int,
                         cell_width_inches: float = 6.0,
                         default_size_pt: int = DEFAULT_FONT_SIZE_PT,
                         min_size_pt: int = MIN_FONT_SIZE_PT,
                         reserved_height_pt: float = 0) -> int:
    """
    Calcula el tamaño de fuente apropiado para que el texto quepa
    dentro del espacio disponible de la celda.

    Reduce progresivamente desde default_size_pt hasta min_size_pt (7pt)
    hasta que el texto estimado quepa en el espacio.
    """
    if not text:
        return default_size_pt

    available_height_pt = available_height_emu / EMU_PER_PT - reserved_height_pt

    for size in range(default_size_pt, min_size_pt - 1, -1):
        # Estimación de ancho promedio de carácter (Arial/Liberation Sans)
        avg_char_width_pt = size * 0.52
        chars_per_line = max(1, int(cell_width_inches * 72 / avg_char_width_pt))
        line_height_pt = size * 1.4
        max_lines = max(1, int(available_height_pt / line_height_pt))

        # Calcular líneas necesarias considerando saltos de línea
        lines_needed = 0
        for line in text.split('\n'):
            if not line.strip():
                lines_needed += 1
            else:
                lines_needed += max(1, -(-len(line) // chars_per_line))

        if lines_needed <= max_lines:
            return size

    return min_size_pt


def _add_auto_sized_text(cell, text: str, available_height_emu: int,
                         cell_width_inches: float = 5.8,
                         reserved_height_pt: float = 0,
                         paragraph_index: int = 0,
                         font_name: str = 'Arial'):
    """
    Agrega texto con auto-sizing a una celda.
    Si el texto es largo, reduce el tamaño de fuente hasta mín 7pt.
    """
    if not text:
        return

    font_size = _calculate_font_size(
        text, available_height_emu,
        cell_width_inches=cell_width_inches,
        reserved_height_pt=reserved_height_pt
    )

    if paragraph_index < len(cell.paragraphs):
        para = cell.paragraphs[paragraph_index]
    else:
        para = cell.add_paragraph()

    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name


# ==================== PARAGRAPH HELPERS ====================

def _remove_extra_paragraphs(cell, keep_count: int = 1):
    """
    Elimina paragraphs extra de una celda, manteniendo solo los primeros `keep_count`.
    Útil para limpiar paragraphs vacíos de la plantilla que desplazan el contenido.
    """
    paragraphs = cell.paragraphs
    for para in paragraphs[keep_count:]:
        p_element = para._element
        p_element.getparent().remove(p_element)


# ==================== CHECKBOX HELPERS ====================

# Contador global para generar nombres únicos de partes de imagen
_checkbox_img_counter = 0


def _create_checked_checkbox_image(original_blob: bytes) -> bytes:
    """
    Dibuja una palomita (✓) encima de la imagen original del checkbox.
    Retorna bytes PNG de la imagen modificada.
    """
    try:
        img = Image.open(BytesIO(original_blob)).convert('RGBA')
    except Exception:
        # Si no se puede abrir (ej: EMF/WMF), crear checkbox desde cero
        img = Image.new('RGBA', (24, 24), (255, 255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.rectangle([0, 0, 23, 23], outline='black', width=1)

    w, h = img.size
    draw = ImageDraw.Draw(img)

    # Dibujar palomita ✓ proporcionalmente
    line_width = max(2, min(w, h) // 6)
    p1 = (int(w * 0.15), int(h * 0.55))
    p2 = (int(w * 0.40), int(h * 0.82))
    p3 = (int(w * 0.85), int(h * 0.18))
    draw.line([p1, p2], fill='black', width=line_width)
    draw.line([p2, p3], fill='black', width=line_width)

    buf = BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def _overlay_checkmark_on_image_run(run, checked: bool):
    """
    Superpone una palomita sobre la imagen de checkbox existente en el run.
    Mantiene la imagen original (recuadro) y dibuja ✓ encima.
    Si checked=False, deja la imagen sin modificar.
    """
    if not checked:
        return

    global _checkbox_img_counter

    # Buscar el blip (referencia a la imagen) dentro del run
    ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    blip = run._element.find(f'.//{{{ns_a}}}blip')
    if blip is None:
        return

    embed_id = blip.get(f'{{{ns_r}}}embed')
    if not embed_id:
        return

    # Obtener la imagen original del paquete DOCX
    part = run.part
    try:
        original_part = part.rels[embed_id].target_part
        original_blob = original_part.blob
    except (KeyError, AttributeError):
        logger.warning(f'No se pudo obtener imagen del checkbox (embed={embed_id})')
        return

    # Crear versión con palomita
    checked_blob = _create_checked_checkbox_image(original_blob)

    # Crear una nueva parte de imagen (para no afectar otros checkboxes que usan la misma imagen)
    _checkbox_img_counter += 1
    new_partname = PackURI(f'/word/media/chk_{_checkbox_img_counter}.png')
    new_image_part = OpcPart(new_partname, 'image/png', checked_blob, part.package)

    # Agregar relación del documento a la nueva imagen
    new_rel_id = part.relate_to(new_image_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

    # Apuntar el blip a la nueva imagen (la original queda intacta para otros checkboxes)
    blip.set(f'{{{ns_r}}}embed', new_rel_id)


# ==================== SOLICITUD DE MANTENIMIENTO ====================

def generate_solicitud_docx(ticket: Ticket) -> BytesIO:
    """
    Genera FORMATO PARA SOLICITUD DE MANTENIMIENTO usando la plantilla original.

    Estructura de la plantilla (según inspección):
    - P[1]: R[0]=img, R[1]=img, R[2]="Folio:                           " (right-aligned, bold, 11pt)
    - Table[0]: 2x2 checkboxes
        Row[0]: "Mantenimiento Preventivo" | vacío
        Row[1]: "Mantenimiento Correctivo" | vacío
    - Table[1]: 1x1 "Área Solicitante: "
    - Table[2]: 4x1
        Row[0]: "Nombre del Solicitante: " + nbsp placeholders
        Row[1]: "Fecha de Elaboración:"
        Row[2]: "Descripción del servicio..." (label)
        Row[3]: vacío (height=2727960 EMU) → descripción con auto-sizing
    """
    template_path = _get_template_path(TEMPLATE_SOLICITUD)
    doc = Document(template_path)

    # 1. Folio - P[1] R[2] contiene "Folio:                           "
    #    Reemplazar espacios trailing con el número de ticket
    if len(doc.paragraphs) > 1 and len(doc.paragraphs[1].runs) > 2:
        folio_run = doc.paragraphs[1].runs[2]
        folio_run.text = 'Folio:  ' + ticket.ticket_number

    # 2. Checkboxes de tipo de mantenimiento - Table[0] Cell[row, 1] (celda vacía)
    table0 = doc.tables[0]
    if ticket.maintenance_type == 'PREVENTIVO':
        run = table0.rows[0].cells[1].paragraphs[0].add_run('X')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    elif ticket.maintenance_type == 'CORRECTIVO':
        run = table0.rows[1].cells[1].paragraphs[0].add_run('X')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    # 3. Área Solicitante - Table[1] Cell[0,0]
    #    Runs existentes: R[0]="Área Solicitante:" R[1]=" " R[2]=" "
    table1 = doc.tables[1]
    dept_name = ticket.requester_department.name if ticket.requester_department else ''
    run = table1.rows[0].cells[0].paragraphs[0].add_run(dept_name)
    run.font.name = 'Arial'

    # 4. Datos del solicitante - Table[2]
    table2 = doc.tables[2]

    # Row[0]: "Nombre del Solicitante: " + \xa0\xa0 (placeholders)
    #    Limpiar placeholders \xa0 y agregar nombre
    cell_nombre = table2.rows[0].cells[0]
    runs_nombre = cell_nombre.paragraphs[0].runs
    if len(runs_nombre) > 1:
        runs_nombre[1].text = ''  # Limpiar \xa0
    if len(runs_nombre) > 2:
        runs_nombre[2].text = ''  # Limpiar \xa0
    req_name = ticket.requester.full_name if ticket.requester else ''
    run = cell_nombre.paragraphs[0].add_run(req_name)
    run.font.name = 'Arial'

    # Row[1]: "Fecha de Elaboración:"
    run = table2.rows[1].cells[0].paragraphs[0].add_run(' ' + _format_date(ticket.created_at))
    run.font.name = 'Arial'

    # Row[3]: Descripción - celda vacía, auto-sizing (height=2727960 EMU ≈ 2.98 inches)
    desc_text = ticket.description or ''
    _add_auto_sized_text(
        table2.rows[3].cells[0],
        desc_text,
        available_height_emu=2727960,
        cell_width_inches=5.8,
        font_name='Arial'
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==================== ORDEN DE TRABAJO ====================

def generate_orden_trabajo_docx(ticket: Ticket) -> BytesIO:
    """
    Genera FORMATO PARA ORDEN DE TRABAJO usando la plantilla original.
    Solo para tickets resueltos o cerrados.

    Estructura de la plantilla (según inspección):
    - Table[0]: 1x2 → "Folio:" | vacío
    - Table[1]: 7x2 (filas 0-4 tienen celdas mergeadas):
        Row[0] h=311150:  Tipo de servicio con checkbox IMG
        Row[1] h=212090:  "Asignado a: "
        Row[2] h=234950:  "Fecha de realización: " + \xa0 placeholders
        Row[3] h=2077085: "TRABAJO REALIZADO:" + 4 paragraphs vacíos
        Row[4] h=901700:  "OBSERVACIÓN:" + 1 paragraph vacío
        Row[5] h=394970:  Cell[0]="Verificado y liberado por:" | Cell[1]="Fecha y firma:"
        Row[6] h=394970:  Cell[0]="Aprobado por:" | Cell[1]="Fecha y firma:"
    """
    if not ticket.is_resolved and ticket.status != 'CLOSED':
        raise ValueError('Solo se puede generar orden de trabajo para tickets resueltos o cerrados')

    template_path = _get_template_path(TEMPLATE_ORDEN)
    doc = Document(template_path)

    # 1. Folio - Table[0] Cell[0,1] (celda vacía a la derecha de "Folio:")
    table0 = doc.tables[0]
    run = table0.rows[0].cells[1].paragraphs[0].add_run(ticket.ticket_number)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # 2. Tabla principal - Table[1]
    table1 = doc.tables[1]

    # Row[0]: Tipo de servicio - superponer ✓ sobre checkboxes marcados
    _mark_service_type_overlay(table1.rows[0].cells[0], ticket)

    # Row[1]: "Asignado a: " - R[0-2]=label bold, R[3]=" "
    #    Agregar nombre después del último run
    assigned_name = ticket.assigned_to.full_name if ticket.assigned_to else ''
    run = table1.rows[1].cells[0].paragraphs[0].add_run(assigned_name)
    run.font.name = 'Arial'

    # Row[2]: "Fecha de realización: " - R[0-4]=label bold, R[5-6]=\xa0 placeholders
    #    Limpiar placeholders y agregar fecha
    cell_fecha = table1.rows[2].cells[0]
    runs_fecha = cell_fecha.paragraphs[0].runs
    if len(runs_fecha) > 5:
        runs_fecha[5].text = ''  # Limpiar \xa0
    if len(runs_fecha) > 6:
        runs_fecha[6].text = ''  # Limpiar \xa0
    run = cell_fecha.paragraphs[0].add_run(_format_date(ticket.resolved_at))
    run.font.name = 'Arial'

    # Row[3]: TRABAJO REALIZADO - P[0]="TRABAJO REALIZADO:", P[1-4]=vacíos
    #    Limpiar paragraphs vacíos extra (P[2], P[3], P[4]) para que no desplacen el texto
    #    Luego escribir en P[1] con auto-sizing
    cell_trabajo = table1.rows[3].cells[0]
    _remove_extra_paragraphs(cell_trabajo, keep_count=2)  # Mantener P[0]=label, P[1]=contenido
    resolution_text = ticket.resolution_notes or ''
    _add_auto_sized_text(
        cell_trabajo,
        resolution_text,
        available_height_emu=2077085,
        cell_width_inches=5.8,
        reserved_height_pt=20,
        paragraph_index=1,
        font_name='Arial'
    )

    # Row[4]: OBSERVACIÓN - P[0]="OBSERVACIÓN:", P[1]=vacío
    #    Auto-sizing en P[1] (height=901700 EMU, reservar ~18pt para label)
    obs_text = ticket.observations or ''
    _add_auto_sized_text(
        table1.rows[4].cells[0],
        obs_text,
        available_height_emu=901700,
        cell_width_inches=5.8,
        reserved_height_pt=18,
        paragraph_index=1,
        font_name='Arial'
    )

    # Row[5]: Verificado y liberado por
    #    Cell[0]: P[0]=vacío, P[1]="Verificado y liberado por: ", P[2]=vacío
    #    → Agregar nombre en P[1] después de "por: "
    requester_name = ticket.requester.full_name if ticket.requester else ''
    cell_verif = table1.rows[5].cells[0]
    run = cell_verif.paragraphs[1].add_run(requester_name)
    run.font.name = 'Arial'

    #    Cell[1]: P[0]=vacío, P[1]="Fecha y firma: "
    #    → Agregar fecha debajo del label
    cell_fecha_verif = table1.rows[5].cells[1]
    run = cell_fecha_verif.paragraphs[1].add_run('\n' + _format_date(ticket.resolved_at))
    run.font.name = 'Arial'
    run.bold = False

    # Row[6]: Aprobado por
    #    Cell[0]: P[0]=vacío, P[1]="Aprobado por: ", P[2]=vacío
    #    → Agregar nombre en P[1] después de "por:"
    resolved_by_name = ticket.resolved_by.full_name if ticket.resolved_by else ''
    cell_aprob = table1.rows[6].cells[0]
    run = cell_aprob.paragraphs[1].add_run(resolved_by_name)
    run.font.name = 'Arial'

    #    Cell[1]: P[0]=vacío, P[1]="Fecha y firma: "
    #    → Agregar fecha debajo del label
    cell_fecha_aprob = table1.rows[6].cells[1]
    run = cell_fecha_aprob.paragraphs[1].add_run('\n' + _format_date(ticket.resolved_at))
    run.font.name = 'Arial'
    run.bold = False

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _mark_service_type_overlay(cell, ticket: Ticket):
    """
    Superpone palomitas (✓) sobre las imágenes de checkbox marcadas.
    Mantiene las imágenes originales (recuadros) intactas.

    Estructura de runs en P[0]:
    R[0-6]: "Tipo de servicio: " (labels en bold, Times New Roman)
    R[7]:   "Mantenimiento:     "
    R[8]:   " Interno"           → service_origin == INTERNO
    R[9]:   [IMG checkbox]       → overlay ✓ si Interno
    R[10]:  "   Preventivo"      → maintenance_type == PREVENTIVO
    R[11]:  [IMG checkbox]       → overlay ✓ si Preventivo
    R[12]:  "   Correctivo"      → maintenance_type == CORRECTIVO
    R[13]:  [IMG checkbox]       → overlay ✓ si Correctivo
    R[14]:  "   Externo"         → service_origin == EXTERNO
    R[15]:  [IMG checkbox]       → overlay ✓ si Externo
    """
    para = cell.paragraphs[0]
    runs = para.runs

    if len(runs) < 16:
        logger.warning(f'Estructura de runs inesperada en tipo de servicio: {len(runs)} runs')
        return

    # Mapeo: índice del IMG run → condición para marcarlo
    checkbox_map = {
        9: ticket.service_origin == 'INTERNO',
        11: ticket.maintenance_type == 'PREVENTIVO',
        13: ticket.maintenance_type == 'CORRECTIVO',
        15: ticket.service_origin == 'EXTERNO',
    }

    for idx, is_checked in checkbox_map.items():
        _overlay_checkmark_on_image_run(runs[idx], is_checked)


# ==================== CONVERSIÓN A PDF ====================

def _find_libreoffice_cmd():
    """Busca el ejecutable de LibreOffice en el sistema."""
    # Linux (Docker)
    for cmd in ['libreoffice', 'soffice']:
        if shutil.which(cmd):
            return cmd

    # Windows (desarrollo local)
    win_paths = [
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ]
    for path in win_paths:
        if os.path.exists(path):
            return path

    return None


def _convert_docx_to_pdf(docx_buffer: BytesIO) -> BytesIO:
    """
    Convierte un DOCX (BytesIO) a PDF usando LibreOffice CLI.
    Retorna BytesIO con el PDF.
    """
    lo_cmd = _find_libreoffice_cmd()
    if not lo_cmd:
        raise RuntimeError(
            'LibreOffice no está instalado. '
            'Instálalo para la conversión a PDF, o descarga en formato DOCX.'
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'document.docx')
        with open(docx_path, 'wb') as f:
            f.write(docx_buffer.read())

        result = subprocess.run(
            [
                lo_cmd, '--headless', '--convert-to', 'pdf',
                '--outdir', tmpdir, docx_path
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            logger.error(f'LibreOffice error: {result.stderr}')
            raise RuntimeError(f'Error al convertir a PDF: {result.stderr}')

        pdf_path = os.path.join(tmpdir, 'document.pdf')
        if not os.path.exists(pdf_path):
            raise RuntimeError('No se generó el archivo PDF')

        pdf_buffer = BytesIO()
        with open(pdf_path, 'rb') as f:
            pdf_buffer.write(f.read())
        pdf_buffer.seek(0)
        return pdf_buffer


# ==================== FUNCIONES PÚBLICAS (PDF) ====================

def generate_solicitud_pdf(ticket: Ticket) -> BytesIO:
    """Genera Solicitud como PDF (DOCX → PDF via LibreOffice)."""
    docx_buffer = generate_solicitud_docx(ticket)
    return _convert_docx_to_pdf(docx_buffer)


def generate_orden_trabajo_pdf(ticket: Ticket) -> BytesIO:
    """Genera Orden de Trabajo como PDF (DOCX → PDF via LibreOffice)."""
    docx_buffer = generate_orden_trabajo_docx(ticket)
    return _convert_docx_to_pdf(docx_buffer)


# ==================== GENERACIÓN EN LOTE ====================

def generate_batch_zip(tickets: List[Ticket], doc_type: str, doc_format: str) -> BytesIO:
    """
    Genera múltiples documentos como ZIP.

    Args:
        tickets: Lista de tickets
        doc_type: 'solicitud', 'orden_trabajo' o 'combinado'
        doc_format: 'pdf' o 'docx'
    """
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for ticket in tickets:
            try:
                if doc_type == 'combinado':
                    _add_combined_to_zip(zf, ticket, doc_format)
                else:
                    _add_single_type_to_zip(zf, ticket, doc_type, doc_format)
            except Exception as e:
                logger.error(f'Error generando documento para {ticket.ticket_number}: {e}')
                continue

    zip_buffer.seek(0)
    return zip_buffer


def _add_single_type_to_zip(zf, ticket: Ticket, doc_type: str, doc_format: str):
    """Agrega un solo tipo de documento al ZIP."""
    generators = {
        ('solicitud', 'pdf'): generate_solicitud_pdf,
        ('solicitud', 'docx'): generate_solicitud_docx,
        ('orden_trabajo', 'pdf'): generate_orden_trabajo_pdf,
        ('orden_trabajo', 'docx'): generate_orden_trabajo_docx,
    }

    gen_func = generators.get((doc_type, doc_format))
    if not gen_func:
        raise ValueError(f'Combinación inválida: {doc_type}/{doc_format}')

    if doc_type == 'orden_trabajo' and not ticket.is_resolved and ticket.status != 'CLOSED':
        logger.warning(f'Omitiendo {ticket.ticket_number}: no resuelto')
        return

    prefix = 'Solicitud' if doc_type == 'solicitud' else 'OrdenTrabajo'
    file_buffer = gen_func(ticket)
    zf.writestr(f'{prefix}_{ticket.ticket_number}.{doc_format}', file_buffer.read())


def _add_combined_to_zip(zf, ticket: Ticket, doc_format: str):
    """Agrega solicitud + orden de trabajo al ZIP para un ticket."""
    # Siempre generar solicitud
    if doc_format == 'pdf':
        sol_buf = generate_solicitud_pdf(ticket)
    else:
        sol_buf = generate_solicitud_docx(ticket)
    zf.writestr(f'Solicitud_{ticket.ticket_number}.{doc_format}', sol_buf.read())

    # Orden solo si está resuelto
    if ticket.is_resolved or ticket.status == 'CLOSED':
        if doc_format == 'pdf':
            ot_buf = generate_orden_trabajo_pdf(ticket)
        else:
            ot_buf = generate_orden_trabajo_docx(ticket)
        zf.writestr(f'OrdenTrabajo_{ticket.ticket_number}.{doc_format}', ot_buf.read())


def generate_concatenated_pdf(tickets: List[Ticket], doc_type: str) -> BytesIO:
    """
    Genera un PDF único con todos los tickets concatenados.
    Usa pypdf para unir los PDFs individuales.
    Soporta doc_type 'combinado' (solicitud + orden por cada ticket).
    """
    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError:
        logger.warning('pypdf no disponible, generando ZIP como fallback')
        return generate_batch_zip(tickets, doc_type, 'pdf')

    writer = PdfWriter()

    for ticket in tickets:
        try:
            if doc_type == 'combinado':
                # Solicitud siempre
                _append_pdf_pages(writer, generate_solicitud_pdf(ticket))
                # Orden solo si resuelto
                if ticket.is_resolved or ticket.status == 'CLOSED':
                    _append_pdf_pages(writer, generate_orden_trabajo_pdf(ticket))
            elif doc_type == 'solicitud':
                _append_pdf_pages(writer, generate_solicitud_pdf(ticket))
            elif doc_type == 'orden_trabajo':
                if not ticket.is_resolved and ticket.status != 'CLOSED':
                    continue
                _append_pdf_pages(writer, generate_orden_trabajo_pdf(ticket))
        except Exception as e:
            logger.error(f'Error concatenando PDF para {ticket.ticket_number}: {e}')
            continue

    if len(writer.pages) == 0:
        raise ValueError('No se generó contenido para ningún ticket')

    buffer = BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer


def _append_pdf_pages(writer, pdf_buffer: BytesIO):
    """Helper para agregar páginas de un PDF al writer."""
    from pypdf import PdfReader
    reader = PdfReader(pdf_buffer)
    for page in reader.pages:
        writer.add_page(page)


# ==================== COMBINADO (TICKET INDIVIDUAL) ====================

def generate_combined_pdf(ticket: Ticket) -> BytesIO:
    """
    Genera un solo PDF con Solicitud + Orden de Trabajo concatenados.
    La solicitud siempre se incluye. La orden solo si el ticket está resuelto.
    """
    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError:
        raise RuntimeError('pypdf requerido para documentos combinados PDF')

    writer = PdfWriter()

    # Siempre incluir solicitud
    _append_pdf_pages(writer, generate_solicitud_pdf(ticket))

    # Orden solo si resuelto
    if ticket.is_resolved or ticket.status == 'CLOSED':
        _append_pdf_pages(writer, generate_orden_trabajo_pdf(ticket))

    buffer = BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer


def generate_combined_docx_zip(ticket: Ticket) -> BytesIO:
    """
    Genera un ZIP con Solicitud.docx + OrdenTrabajo.docx para un ticket.
    (No es posible fusionar dos plantillas Word diferentes en un solo DOCX
    sin perder formato, así que se entregan en ZIP)
    """
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        sol_buf = generate_solicitud_docx(ticket)
        zf.writestr(f'Solicitud_{ticket.ticket_number}.docx', sol_buf.read())

        if ticket.is_resolved or ticket.status == 'CLOSED':
            ot_buf = generate_orden_trabajo_docx(ticket)
            zf.writestr(f'OrdenTrabajo_{ticket.ticket_number}.docx', ot_buf.read())

    zip_buffer.seek(0)
    return zip_buffer
